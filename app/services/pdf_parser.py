"""Document parsing service for extracting learning units from PDF and Word files."""

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import pdfplumber

try:
    from docx import Document as WordDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.config import settings
from app.models.learning_unit import UnitType
from app.schemas.learning_unit import LearningUnitCreate

logger = logging.getLogger(__name__)


# Polish + Latin letters used for word-boundary checks
_PL_LETTERS = r"a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ"


def _make_word_boundary_pattern(word: str) -> re.Pattern:
    """Return a case-insensitive pattern that matches `word` only at word boundaries.

    Standard ``\\b`` does not work with Polish diacritics in Python's ``re``
    module, so we use manual negative look-around for letter characters instead.
    """
    return re.compile(
        r"(?<![" + _PL_LETTERS + r"])" + re.escape(word) + r"(?![" + _PL_LETTERS + r"])",
        re.IGNORECASE,
    )


def extract_sentence_for_word(word: str, full_text: str) -> Optional[str]:
    if not word or not full_text:
        return None

    # Split on any sentence boundary: after . ! ? followed by whitespace.
    # The previous pattern required an uppercase letter after the boundary,
    # which missed sentences starting with lowercase (common in PDF extracts).
    sentences = re.split(r"(?<=[.!?])\s+", full_text)

    word_lower = word.lower().strip()
    word_parts = word_lower.split()
    candidates: list[str] = []

    if len(word_parts) == 1:
        # Single word: require word-boundary match to avoid false positives
        # (e.g. "ma" must not match inside "mama").
        pattern = _make_word_boundary_pattern(word_lower)
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 5 or len(sentence) > 250:
                continue
            if pattern.search(sentence):
                candidates.append(sentence)
    else:
        # Multi-word phrase: substring match is acceptable because the phrase
        # itself acts as a natural boundary.
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 5 or len(sentence) > 250:
                continue
            if word_lower in sentence.lower():
                candidates.append(sentence)

    if not candidates:
        return None

    return min(candidates, key=len)


@dataclass
class ParseResult:
    """Result of parsing a single line."""
    text: str
    translation: str
    unit_type: UnitType
    part_of_speech: Optional[str] = None


@dataclass
class PDFParseResult:
    """Result of parsing an entire PDF."""
    units: list[LearningUnitCreate]
    skipped_lines: int
    total_lines: int


class PDFParser:
    """
    Parser for extracting learning units from PDF files.
    
    Handles line-by-line parsing with delimiter detection,
    unit type classification, and part of speech extraction.
    """
    
    # Delimiters in order of priority (longer first)
    DELIMITERS = settings.pdf_delimiters
    
    # Regex for extracting part of speech from parentheses
    # Matches: "word (noun)" -> groups: ("word", "noun")
    POS_PATTERN = re.compile(r"^(.+?)\s*\(([^)]+)\)\s*$")
    
    # Sentence-ending punctuation
    SENTENCE_ENDINGS = {".", "!", "?"}
    
    # Pattern to detect repeated characters (3+ consecutive same char)
    REPEATED_CHAR_PATTERN = re.compile(r"(.)\1{2,}")
    
    def __init__(self, source_filename: str):
        """
        Initialize parser with source filename.
        
        Args:
            source_filename: Name of the PDF file being parsed.
        """
        self.source_filename = source_filename
    
    def parse_file(self, file_path: Path) -> PDFParseResult:
        """
        Parse a document file (PDF or DOCX) and extract all learning units.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            PDFParseResult with extracted units and statistics.
        """
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def parse_pdf(self, pdf_path: Path) -> PDFParseResult:
        """Legacy method for backwards compatibility."""
        return self._parse_pdf(pdf_path)
    
    def _parse_pdf(self, pdf_path: Path) -> PDFParseResult:
        """
        Parse a PDF file and extract all learning units.
        
        Args:
            pdf_path: Path to the PDF file.
            
        Returns:
            PDFParseResult with extracted units and statistics.
        """
        units: list[LearningUnitCreate] = []
        skipped_lines = 0
        total_lines = 0

        with pdfplumber.open(pdf_path) as pdf:
            full_text_parts: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                full_text_parts.append(page_text)

            full_text = "\n".join(full_text_parts)

            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if not page_text:
                    continue

                lines = page_text.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    total_lines += 1
                    
                    try:
                        result = self._parse_line(line)
                        if result:
                            context_sentence = extract_sentence_for_word(result.text, full_text)
                            unit = LearningUnitCreate(
                                text=result.text,
                                type=result.unit_type,
                                part_of_speech=result.part_of_speech,
                                translation=result.translation,
                                context_sentence=context_sentence,
                                source_pdf=self.source_filename,
                                page_number=page_num,
                            )
                            units.append(unit)
                        else:
                            skipped_lines += 1
                            logger.warning(
                                f"Skipped line (no delimiter found): '{line[:50]}...'"
                            )
                    except Exception as e:
                        skipped_lines += 1
                        logger.warning(
                            f"Failed to parse line: '{line[:50]}...' - Error: {e}"
                        )

        with_context = sum(1 for u in units if u.context_sentence)
        logger.info(f"Context sentences extracted: {with_context}/{len(units)}")

        return PDFParseResult(
            units=units,
            skipped_lines=skipped_lines,
            total_lines=total_lines,
        )
    
    def _extract_cell_text(self, cell) -> str:
        """
        Extract all text from a table cell, handling nested paragraphs.
        
        Args:
            cell: python-docx table cell object.
            
        Returns:
            Complete text content of the cell.
        """
        # Try the simple approach first
        text = cell.text.strip()
        
        # If that doesn't work or seems incomplete, extract from paragraphs
        if not text or len(text) < 2:
            # Extract from all paragraphs in the cell
            parts = []
            for para in cell.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    parts.append(para_text)
            text = "\n".join(parts).strip()
        
        return text
    
    def _parse_docx(self, docx_path: Path) -> PDFParseResult:
        """
        Parse a Word document (.docx) and extract all learning units.
        
        Handles both paragraphs and tables. Tables are parsed assuming:
        - Two-column format: column 0 = source text, column 1 = translation
        - Or single-column format with delimiters (same as paragraphs)
        
        Args:
            docx_path: Path to the Word file.
            
        Returns:
            PDFParseResult with extracted units and statistics.
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
        
        units: list[LearningUnitCreate] = []
        skipped_lines = 0
        total_lines = 0
        
        doc = WordDocument(docx_path)

        # Collect paragraph text for a "full_text" used in sentence extraction.
        # Note: table-based vocabularies are structured pairs and intentionally skip extraction.
        para_texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        full_text = "\n".join(para_texts)
        
        # Parse paragraphs (line-by-line format)
        for para in doc.paragraphs:
            line = para.text.strip()
            if not line:
                continue
            
            total_lines += 1
            
            try:
                result = self._parse_line(line)
                if result:
                    context_sentence = extract_sentence_for_word(result.text, full_text)
                    unit = LearningUnitCreate(
                        text=result.text,
                        type=result.unit_type,
                        part_of_speech=result.part_of_speech,
                        translation=result.translation,
                        context_sentence=context_sentence,
                        source_pdf=self.source_filename,
                        page_number=None,  # Word docs don't have page numbers easily
                    )
                    units.append(unit)
                else:
                    skipped_lines += 1
                    logger.debug(
                        f"Skipped paragraph (no delimiter found): '{line[:50]}...'"
                    )
            except Exception as e:
                skipped_lines += 1
                logger.warning(
                    f"Failed to parse paragraph: '{line[:50]}...' - Error: {e}"
                )
        
        # Parse tables (two-column format common in vocabulary files)
        table_count = 0
        for table_idx, table in enumerate(doc.tables):
            table_count += 1
            row_count = 0
            rows_processed = 0
            
            logger.info(f"Processing table {table_idx}: {len(table.rows)} rows, {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                row_count += 1
                # Use improved cell text extraction
                cells = [self._extract_cell_text(cell) for cell in row.cells]
                cells = [c for c in cells if c]  # Filter out empty cells
                
                # Skip completely empty rows
                if not any(cells):
                    continue
                
                # Handle multi-column format (use first 2 columns)
                if len(cells) >= 2:
                    # Two-column format: assume first column = source, second = translation
                    # Also try if cells are reversed (second might be source, first translation)
                    source_cell = cells[0].strip() if cells[0] else ""
                    translation_cell = cells[1].strip() if cells[1] else ""
                    
                    # Log raw cell contents for debugging (first few rows only to avoid spam)
                    if row_idx <= 3:
                        logger.info(
                            f"Table {table_idx}, row {row_idx}: Raw cells = {[c[:40] + '...' if len(c) > 40 else c for c in cells]}"
                        )
                    
                    # Check if cells contain multiple entries (separated by newlines)
                    source_lines = [line.strip() for line in source_cell.split("\n") if line.strip()]
                    translation_lines = [line.strip() for line in translation_cell.split("\n") if line.strip()]
                    
                    # If both cells have the same number of lines, treat as paired entries
                    if len(source_lines) == len(translation_lines) and len(source_lines) > 1:
                        # Multiple entries in this row
                        for src, trans in zip(source_lines, translation_lines):
                            if not src or not trans:
                                continue
                            total_lines += 1
                            try:
                                part_of_speech = None
                                pos_match = self.POS_PATTERN.match(src)
                                if pos_match:
                                    src = pos_match.group(1).strip()
                                    part_of_speech = pos_match.group(2).strip()
                                
                                unit_type = self._detect_unit_type(src)
                                unit = LearningUnitCreate(
                                    text=src,
                                    type=unit_type,
                                    part_of_speech=part_of_speech,
                                    translation=trans,
                                    context_sentence=None,
                                    source_pdf=self.source_filename,
                                    page_number=None,
                                )
                                units.append(unit)
                                rows_processed += 1
                            except Exception as e:
                                skipped_lines += 1
                                logger.warning(f"Table {table_idx}, row {row_idx}: Failed to parse multi-entry: {e}")
                        continue
                    
                    # Single entry per row (normal case)
                    source_text = source_cell
                    translation = translation_cell
                    
                    # Skip if either is empty (might be header row or empty row)
                    if not source_text or not translation:
                        logger.debug(
                            f"Table {table_idx}, row {row_idx}: Skipping row with empty cell(s) - "
                            f"'{source_text[:30] if source_text else '(empty)'}' / "
                            f"'{translation[:30] if translation else '(empty)'}'"
                        )
                        skipped_lines += 1
                        continue
                    
                    total_lines += 1
                    
                    try:
                        # Clean up any trailing whitespace or formatting
                        source_text = source_text.strip()
                        translation = translation.strip()
                        
                        # Extract part of speech if present in source text
                        part_of_speech = None
                        pos_match = self.POS_PATTERN.match(source_text)
                        if pos_match:
                            source_text = pos_match.group(1).strip()
                            part_of_speech = pos_match.group(2).strip()
                        
                        # Detect unit type
                        unit_type = self._detect_unit_type(source_text)
                        
                        unit = LearningUnitCreate(
                            text=source_text,
                            type=unit_type,
                            part_of_speech=part_of_speech,
                            translation=translation,
                            context_sentence=None,
                            source_pdf=self.source_filename,
                            page_number=None,
                        )
                        units.append(unit)
                        rows_processed += 1
                        logger.debug(
                            f"Table {table_idx}, row {row_idx}: Extracted '{source_text[:30]}...' = '{translation[:30]}...'"
                        )
                    except Exception as e:
                        skipped_lines += 1
                        logger.warning(
                            f"Table {table_idx}, row {row_idx}: Failed to parse - "
                            f"'{cells[0][:30] if cells[0] else '(empty)'}...' / "
                            f"'{cells[1][:30] if len(cells) > 1 and cells[1] else '(empty)'}...' - Error: {e}"
                        )
                elif len(cells) == 1:
                    # Single-column format: try to parse with delimiter
                    line = cells[0].strip()
                    if not line:
                        continue
                    
                    total_lines += 1
                    
                    try:
                        result = self._parse_line(line)
                        if result:
                            unit = LearningUnitCreate(
                                text=result.text,
                                type=result.unit_type,
                                part_of_speech=result.part_of_speech,
                                translation=result.translation,
                                context_sentence=None,
                                source_pdf=self.source_filename,
                                page_number=None,
                            )
                            units.append(unit)
                            rows_processed += 1
                        else:
                            skipped_lines += 1
                            logger.debug(
                                f"Table {table_idx}, row {row_idx}: Skipped (no delimiter found): '{line[:50]}...'"
                            )
                    except Exception as e:
                        skipped_lines += 1
                        logger.warning(
                            f"Table {table_idx}, row {row_idx}: Failed to parse cell: '{line[:50]}...' - Error: {e}"
                        )
                else:
                    # Row has no cells or all empty
                    skipped_lines += 1
                    logger.debug(f"Table {table_idx}, row {row_idx}: Skipped (no valid cells)")
            
            logger.info(
                f"Table {table_idx}: Processed {rows_processed} units from {row_count} rows "
                f"({row_count - rows_processed} skipped)"
            )
        
        logger.info(
            f"Parsed DOCX {docx_path.name}: {len(units)} units from {total_lines} lines/cells "
            f"({len(doc.paragraphs)} paragraphs, {table_count} tables with {sum(len(t.rows) for t in doc.tables)} total rows, {skipped_lines} skipped)"
        )
        
        # Verbose summary for debugging
        if table_count > 0:
            for table_idx, table in enumerate(doc.tables):
                non_empty_rows = sum(
                    1 for row in table.rows 
                    if any(self._extract_cell_text(cell).strip() for cell in row.cells)
                )
                logger.info(
                    f"  Table {table_idx}: {len(table.rows)} rows ({non_empty_rows} non-empty), "
                    f"{len(table.columns)} columns"
                )

        with_context = sum(1 for u in units if u.context_sentence)
        logger.info(f"Context sentences extracted: {with_context}/{len(units)}")

        return PDFParseResult(
            units=units,
            skipped_lines=skipped_lines,
            total_lines=total_lines,
        )
    
    def _parse_line(self, line: str) -> Optional[ParseResult]:
        """
        Parse a single line into source text and translation.
        
        Args:
            line: Raw line from PDF.
            
        Returns:
            ParseResult if successful, None if line cannot be parsed.
        """
        # Clean up duplicated characters from PDF extraction issues
        line = self._deduplicate_chars(line)
        
        # Try each delimiter in order
        source_text = None
        translation = None
        
        for delimiter in self.DELIMITERS:
            if delimiter in line:
                parts = line.split(delimiter, 1)
                if len(parts) == 2:
                    source_text = parts[0].strip()
                    translation = parts[1].strip()
                    break
        
        if not source_text or not translation:
            return None
        
        # Extract part of speech if present
        part_of_speech = None
        pos_match = self.POS_PATTERN.match(source_text)
        if pos_match:
            source_text = pos_match.group(1).strip()
            part_of_speech = pos_match.group(2).strip()
        
        # Detect unit type
        unit_type = self._detect_unit_type(source_text)
        
        return ParseResult(
            text=source_text,
            translation=translation,
            unit_type=unit_type,
            part_of_speech=part_of_speech,
        )
    
    def _detect_unit_type(self, text: str) -> UnitType:
        """
        Detect the type of learning unit based on content.
        
        Rules:
        - Ends with . ! ? → sentence
        - Contains spaces but no sentence ending → phrase
        - No spaces and no sentence ending → word
        
        Args:
            text: Source text to classify.
            
        Returns:
            UnitType classification.
        """
        text = text.strip()
        
        # Check for sentence-ending punctuation
        if text and text[-1] in self.SENTENCE_ENDINGS:
            return UnitType.SENTENCE
        
        # Check for spaces (indicating multiple words)
        if " " in text:
            return UnitType.PHRASE
        
        return UnitType.WORD
    
    def _deduplicate_chars(self, text: str) -> str:
        """
        Remove consecutive duplicate characters caused by PDF text layer issues.
        
        Some PDFs have overlapping text layers that cause each character
        to be extracted multiple times (e.g., "łłłłaaaaggggooooddddnnnnyyyy" → "łagodny").
        
        This detects patterns where the same character repeats 3+ times
        and reduces them to a single character.
        
        Args:
            text: Raw text from PDF.
            
        Returns:
            Cleaned text with duplicates removed.
        """
        # Check if text has suspicious repeated patterns (same char 3+ times)
        if not self.REPEATED_CHAR_PATTERN.search(text):
            return text  # No repeated chars, return as-is
        
        # Detect the repetition factor by analyzing the pattern
        # Look at first few chars to determine repetition count
        if len(text) < 2:
            return text
        
        # Find the most common repetition count
        matches = self.REPEATED_CHAR_PATTERN.findall(text)
        if not matches:
            return text
        
        # Count consecutive occurrences to detect the pattern
        i = 0
        rep_counts = []
        while i < len(text):
            char = text[i]
            count = 1
            while i + count < len(text) and text[i + count] == char:
                count += 1
            if count > 1:
                rep_counts.append(count)
            i += count
        
        if not rep_counts:
            return text
        
        # Find the most common repetition factor
        from collections import Counter
        count_freq = Counter(rep_counts)
        most_common_rep = count_freq.most_common(1)[0][0]
        
        # Only deduplicate if there's a consistent pattern (most chars repeat same amount)
        # and the repetition is >= 2
        if most_common_rep < 2:
            return text
        
        # Check if this pattern is consistent across most of the text
        consistent_count = sum(1 for c in rep_counts if c == most_common_rep)
        if consistent_count < len(rep_counts) * 0.5:
            return text  # Pattern not consistent enough
        
        # Deduplicate: reduce each repeated char group to single char
        result = []
        i = 0
        while i < len(text):
            char = text[i]
            count = 1
            while i + count < len(text) and text[i + count] == char:
                count += 1
            result.append(char)
            i += count
        
        cleaned = "".join(result)
        
        if cleaned != text:
            logger.debug(f"Deduplicated: '{text[:50]}...' → '{cleaned[:50]}...'")
        
        return cleaned
    
    @classmethod
    def parse_line_standalone(cls, line: str) -> Optional[ParseResult]:
        """
        Parse a single line without PDF context (for testing).
        
        Args:
            line: Raw line to parse.
            
        Returns:
            ParseResult if successful, None otherwise.
        """
        parser = cls("test.pdf")
        return parser._parse_line(line)
